import re
import io
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from docx import Document
import json


def extract_text_from_pdf(path_or_file):
    output_string = io.StringIO()
    laparams = LAParams()
    # pdfminer can accept file-like objects; open path if a string
    if isinstance(path_or_file, str):
        with open(path_or_file, 'rb') as f:
            extract_text_to_fp(f, output_string, laparams=laparams)
    else:
        extract_text_to_fp(path_or_file, output_string, laparams=laparams)
    return output_string.getvalue()


def extract_text_from_docx(path_or_file):
    # python-docx expects a path or a file-like object
    if isinstance(path_or_file, str):
        doc = Document(path_or_file)
    else:
        doc = Document(path_or_file)
    texts = [p.text for p in doc.paragraphs]
    return '\n'.join(texts)


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"\+?\d[\d\s().-]{6,}\d")
NAME_RE = re.compile(r"^(?:Name|Full Name)[:\s]*([A-Za-z ,.'-]{2,})$", re.IGNORECASE | re.MULTILINE)


def find_emails(text):
    return list(set(EMAIL_RE.findall(text)))


def find_phones(text):
    phones = PHONE_RE.findall(text)
    # quick cleanup
    cleaned = [re.sub(r"[\s().-]", "", p) for p in phones]
    return list(set(cleaned))


def find_name(text):
    m = NAME_RE.search(text)
    if m:
        return m.group(1).strip()
    # fallback: first non-empty line with letters
    for line in text.splitlines():
        line = line.strip()
        if line and len(line.split()) <= 5 and any(c.isalpha() for c in line):
            # avoid lines that contain email or phone
            if EMAIL_RE.search(line) or PHONE_RE.search(line):
                continue
            return line
    return None


def find_skills(text):
    # naive skills list; extend as needed
    candidates = [
        'Python', 'Java', 'C++', 'C#', 'JavaScript', 'TypeScript', 'SQL', 'NoSQL',
        'Django', 'Flask', 'React', 'Angular', 'Node', 'TensorFlow', 'PyTorch',
        'AWS', 'GCP', 'Azure', 'Docker', 'Kubernetes', 'Git'
    ]
    found = []
    for s in candidates:
        if re.search(r"\b" + re.escape(s) + r"\b", text, re.IGNORECASE):
            found.append(s)
    return found


def parse_resume_text(text):
    data = {}
    data['name'] = find_name(text)
    data['emails'] = find_emails(text)
    data['phones'] = find_phones(text)
    data['skills'] = find_skills(text)

    # very naive sections extraction
    lines = text.splitlines()
    education = []
    experience = []
    current = None
    current_buf = []
    for line in lines:
        l = line.strip()
        if not l:
            continue
        if re.search(r'education', l, re.IGNORECASE):
            if current == 'experience' and current_buf:
                experience.append('\n'.join(current_buf))
            current = 'education'
            current_buf = []
            continue
        if re.search(r'experience|employment|work history', l, re.IGNORECASE):
            if current == 'education' and current_buf:
                education.append('\n'.join(current_buf))
            current = 'experience'
            current_buf = []
            continue
        if current:
            current_buf.append(l)

    if current == 'education' and current_buf:
        education.append('\n'.join(current_buf))
    if current == 'experience' and current_buf:
        experience.append('\n'.join(current_buf))

    data['education'] = education
    data['experience'] = experience

    return data


def parse_file(file_path, content_type=None, file_obj=None):
    # file_obj if provided should be file-like binary
    text = ''
    if file_path.lower().endswith('.pdf') or (content_type and 'pdf' in content_type):
        if file_obj:
            text = extract_text_from_pdf(file_obj)
        else:
            text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith('.docx') or (content_type and 'word' in content_type) or file_path.lower().endswith('.doc'):
        if file_obj:
            # python-docx can accept file-like objects
            text = extract_text_from_docx(file_obj)
        else:
            text = extract_text_from_docx(file_path)
    else:
        # try reading as text
        try:
            if file_obj:
                text = file_obj.read().decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
        except Exception:
            text = ''

    parsed = parse_resume_text(text)
    return parsed


def to_json(data):
    return json.dumps(data, indent=2)
